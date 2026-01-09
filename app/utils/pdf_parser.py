# -----------------------------------------------------------------------------
# app/utils/pdf_parser.py  (now handles PDF and DOCX)
# -----------------------------------------------------------------------------
# install: pip install pymupdf python-docx
import io
import logging
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
import re

logger = logging.getLogger(__name__)

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None  # optional dependency


def _clean_text(text: str) -> str:
    """
    Common text cleanup for extracted document text.
    """
    if not text:
        return ""

    # remove null characters
    text = text.replace("\x00", "")

    # normalize spaces/tabs
    text = re.sub(r"[ \t]+", " ", text)

    # normalize excessive newlines (keep paragraphs)
    text = re.sub(r"\n\s*\n+", "\n\n", text)

    # remove too many consecutive newlines
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def _extract_text_from_doc(doc: fitz.Document) -> str:
    """
    Core extraction logic from a PyMuPDF Document.
    """
    pages = []

    page_count = doc.page_count if hasattr(doc, "page_count") else len(doc)

    for page_num in range(page_count):
        try:
            page = doc.load_page(page_num)

            # Primary extraction
            page_text = page.get_text("text") or ""

            # Fallback if empty (block-based extraction)
            if not page_text.strip():
                blocks = page.get_text("blocks") or []
                page_text = "\n".join(
                    block[4].strip()
                    for block in blocks
                    if len(block) > 4 and isinstance(block[4], str) and block[4].strip()
                )

            page_text = _clean_text(page_text)
            if page_text:
                pages.append(page_text)

        except Exception:
            # Skip problematic pages instead of crashing
            logger.exception("Error extracting page %s", page_num)
            continue

    return "\n\n".join(pages)


def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    """
    Extract text from PDF bytes using PyMuPDF.

    Used when PDF is uploaded directly.
    """
    if not pdf_bytes:
        return ""

    try:
        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            # Handle encrypted PDFs
            if getattr(doc, "is_encrypted", False):
                try:
                    if not doc.authenticate(""):
                        logger.warning("PDF is encrypted and requires a password.")
                        return ""
                except Exception:
                    logger.exception("Failed to authenticate encrypted PDF bytes")
                    return ""

            return _extract_text_from_doc(doc)

    except Exception:
        logger.exception("Failed to open/parse PDF from bytes")
        return ""


def extract_text_from_pdf_path(file_path: str) -> str:
    """
    Extract text from a local PDF file path using PyMuPDF.

    Used when PDF is already saved on disk.
    """
    if not file_path:
        return ""

    path = Path(file_path)
    if not path.is_file():
        logger.warning("PDF path not found: %s", file_path)
        return ""

    try:
        with fitz.open(str(path)) as doc:
            # Handle encrypted PDFs
            if getattr(doc, "is_encrypted", False):
                try:
                    if not doc.authenticate(""):
                        logger.warning("PDF is encrypted and requires a password: %s", file_path)
                        return ""
                except Exception:
                    logger.exception("Failed to authenticate encrypted PDF: %s", file_path)
                    return ""

            return _extract_text_from_doc(doc)

    except Exception:
        logger.exception("Failed to open/parse PDF: %s", file_path)
        return ""


# ---------------- DOCX handling ----------------

def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    """
    Extract text from DOCX bytes using python-docx.
    """
    if not docx_bytes:
        return ""

    if DocxDocument is None:
        logger.warning("python-docx not installed; cannot extract DOCX content from bytes.")
        return ""

    try:
        stream = io.BytesIO(docx_bytes)
        doc = DocxDocument(stream)
        paragraphs = [p.text for p in doc.paragraphs if p.text]

        # also extract simple table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)

        text = "\n\n".join(paragraphs)
        return _clean_text(text)
    except Exception:
        logger.exception("Failed to extract DOCX from bytes")
        return ""


def extract_text_from_docx_path(file_path: str) -> str:
    """
    Extract text from a DOCX file path using python-docx.
    """
    if not file_path:
        return ""

    path = Path(file_path)
    if not path.is_file():
        logger.warning("DOCX path not found: %s", file_path)
        return ""

    if DocxDocument is None:
        logger.warning("python-docx not installed; cannot extract DOCX content: %s", file_path)
        return ""

    try:
        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text]

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)

        text = "\n\n".join(paragraphs)
        return _clean_text(text)
    except Exception:
        logger.exception("Failed to extract DOCX: %s", file_path)
        return ""


# ---------------- Unified helpers ----------------

def extract_text_from_file_path(file_path: str) -> str:
    """
    Extract text from a file (PDF, DOCX, or DOC) by inspecting its extension.
    Falls back to PDF extraction for .pdf and DOCX extraction for .docx.
    For legacy .doc files, attempts to use textract if available.
    """
    if not file_path:
        return ""

    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf_path(file_path)
    elif ext == ".docx":
        return extract_text_from_docx_path(file_path)
    elif ext == ".doc":
        # Try textract if available (supports older .doc)
        try:
            import textract

            text = textract.process(file_path)
            if isinstance(text, bytes):
                text = text.decode("utf-8", errors="ignore")
            return _clean_text(text)
        except Exception:
            logger.exception("Failed to extract .doc using textract. Consider converting to .docx")
            return ""
    else:
        logger.warning("Unsupported file extension for text extraction: %s", ext)
        return ""


def extract_text_from_file_bytes(file_bytes: bytes, filename: Optional[str] = None) -> str:
    """
    Extract text from file bytes (PDF or DOCX). If filename is provided, use its suffix to choose.
    """
    if not file_bytes:
        return ""

    ext = None
    if filename:
        ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return extract_text_from_pdf_bytes(file_bytes)
    elif ext == ".docx":
        return extract_text_from_docx_bytes(file_bytes)
    else:
        # Try PDF extraction first (most common), then DOCX
        txt = extract_text_from_pdf_bytes(file_bytes)
        if txt and txt.strip():
            return txt
        if DocxDocument is not None:
            return extract_text_from_docx_bytes(file_bytes)
        return ""
